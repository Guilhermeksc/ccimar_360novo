from PyQt6.QtWidgets import (
    QFrame, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QLineEdit, QTextEdit,
    QMessageBox, QCheckBox, QGroupBox
)
import sqlite3
import logging
from PyQt6.QtCore import Qt, pyqtSignal
import re
from utils.add_button import add_button_func
from .chatbot_utils.flow_layout import FlowLayout
import ollama

def create_chatbot_local(title_text, database_model, icons):
    """Cria e retorna a interface do chatbot local usando Ollama."""
    content_frame = QFrame()
    layout = QVBoxLayout(content_frame)
    layout.setContentsMargins(10, 10, 10, 10)
    layout.setSpacing(10)

    chatbot_widget = ChatbotWidget(title_text, database_model, icons)
    layout.addWidget(chatbot_widget)

    return content_frame

class ChatbotWidget(QFrame):
    def __init__(self, title_text, database_model, icons, parent=None):
        super().__init__(parent)
        self.icons = icons
        # Verifica se o objeto possui o atributo "database_manager"
        if hasattr(database_model, "database_manager"):
            self.db_path = database_model.database_manager.db_path
        else:
            self.db_path = str(database_model)
        self.db_metadata = self.get_database_metadata()
        self.last_sql_query = None  # Armazena a última consulta SQL
        
        self.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #FFFFFF;
            }
            QLineEdit, QTextEdit {
                background-color: #FFFFFF;
                border-radius: 5px;
                padding: 5px;
                font-size: 14px;
            }
            QPushButton {
                background-color: #6272A4;
                color: #FFF;
                font-weight: bold;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #50FA7B;
            }
            QGroupBox {
                color: white;
                font-size: 14px;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)

        title = QLabel(title_text)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        self.input_field = QLineEdit()
        self.input_field.setPlaceholderText("Faça a sua pergunta...")
        layout.addWidget(self.input_field)

        # Área de seleção para as tabelas
        self.table_selection_group = QGroupBox("Selecionar Tabelas para Análise")
        self.table_selection_layout = QVBoxLayout()
        self.table_checkboxes = {}

        for table_name in self.db_metadata.keys():
            checkbox = QCheckBox(table_name)
            checkbox.setStyleSheet("color: white; font-size: 14px;")
            checkbox.setChecked(False)
            self.table_checkboxes[table_name] = checkbox
            self.table_selection_layout.addWidget(checkbox)

        self.table_selection_group.setLayout(self.table_selection_layout)
        layout.addWidget(self.table_selection_group)

        # Layout horizontal para os botões
        button_layout = QHBoxLayout()

        self.api_button = QPushButton("Pergunte")
        self.api_button.clicked.connect(self.generate_local_response)
        button_layout.addWidget(self.api_button)

        layout.addLayout(button_layout)

        self.response_output = QTextEdit()
        self.response_output.setReadOnly(True)
        layout.addWidget(self.response_output)
    

    def get_database_metadata(self):
        """Lê as tabelas e suas colunas no banco de dados."""
        metadata = {}
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                for table_name in tables:
                    table_name = table_name[0]
                    cursor.execute(f"PRAGMA table_info({table_name})")
                    columns = cursor.fetchall()
                    # Armazena os nomes das colunas com seus tipos, ex.: "cod_siafi (INTEGER)"
                    metadata[table_name] = [f"{col[1]} ({col[2]})" for col in columns]
        except sqlite3.Error as e:
            logging.error(f"Database error: {e}")
            return {}
        return metadata

    def generate_local_response(self):
        """
        Envia a pergunta ao Ollama com metadados das tabelas selecionadas.
        A resposta deve ser uma consulta SQL contendo SEMPRE a coluna 'cod_siafi'.
        """
        user_question = self.input_field.text().strip()
        if not user_question:
            QMessageBox.warning(self, "Aviso", "Por favor, insira uma pergunta.")
            return
        
        # Coletar tabelas e colunas selecionadas
        selected_tables = []
        for table, checkbox in self.table_checkboxes.items():
            if checkbox.isChecked():
                columns_info = ", ".join(self.db_metadata.get(table, []))
                selected_tables.append(f"Tabela: {table}\nColunas: {columns_info}")
        
        if not selected_tables:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione ao menos uma tabela para análise.")
            return
        
        tables_info = "\n\n".join(selected_tables)
        print(tables_info)

        # Definição do prompt com exemplos concretos de SQL para melhorar a resposta
        system_prompt = (
            "Você é um assistente especializado em consultas SQL para SQLite.\n"
            "Seu objetivo é gerar consultas SQL precisas para o banco de dados, seguindo sempre estas regras:\n\n"
            "⚠️ **REGRAS IMPORTANTES** ⚠️\n"
            "- SEMPRE inclua a coluna 'cod_siafi'.\n"
            "- Use `ORDER BY coluna DESC LIMIT 5` para os maiores valores.\n"
            "- Para encontrar o maior valor de uma coluna, use `MAX(coluna)`.\n"
            "- NÃO utilize `GROUP BY` sem necessidade.\n"
            "- Retorne apenas a consulta SQL dentro de ```sql ... ```.\n\n"
            "**Exemplos de consultas corretas:**\n"
            "🔹 **Quais os 5 maiores valores de dispensa?**\n"
            "```sql\n"
            "SELECT cod_siafi, valor_dispensa FROM criterio_execucao_licitacao ORDER BY valor_dispensa DESC LIMIT 5;\n"
            "```\n"
            "🔹 **Qual o maior valor de credenciamento?**\n"
            "```sql\n"
            "SELECT cod_siafi, MAX(valor_credenciamento) AS maior_valor FROM criterio_execucao_licitacao;\n"
            "```\n"
            "\nAgora gere uma consulta SQL para a seguinte pergunta:\n"
            f"**Pergunta do usuário:** {user_question}\n\n"
            f"📂 **Estrutura do Banco de Dados:**\n{tables_info}"
        )

        try:
            response = ollama.chat(
                model="phi3:mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Pergunta: {user_question}"}
                ]
            )

            ai_response = response["message"]["content"]
        except Exception as e:
            logging.error(f"Erro ao conectar ao Ollama: {e}")
            self.response_output.setText("Erro ao acessar o modelo local.")
            return

        output_text = f"🔍 **Consulta SQL gerada:**\n```sql\n{ai_response}\n```\n\n"

        # Extração do SQL da resposta
        sql_query = ""
        if "```" in ai_response:
            parts = ai_response.split("```")
            if len(parts) >= 2:
                sql_query = parts[1].strip()
                if sql_query.lower().startswith("sql"):
                    sql_query = sql_query[3:].strip()
        else:
            sql_query = ai_response.strip()

        # Executa a consulta SQL gerada e exibe os resultados
        if sql_query:
            self.last_sql_query = sql_query  # Armazena a consulta para exportação
            logging.info(f"Consulta SQL armazenada: {self.last_sql_query}")
            result_text = self.execute_sql_query(sql_query)
            output_text += f"📊 **Resultados da Consulta:**\n{result_text}"
        else:
            output_text += "⚠️ Não foi possível extrair a consulta SQL."

        self.response_output.setText(output_text)


    def execute_sql_query(self, query):
        """
        Executa a consulta SQL no banco de dados e retorna os resultados formatados.
        Se a consulta retornar a coluna 'cod_siafi', realiza uma busca adicional na tabela
        'organizacoes_militares' para obter os dados correspondentes.
        """
        def format_currency(value):
            # Formata valores numéricos para o padrão BRL: 1234.56 -> "R$ 1.234,56"
            formatted = f"{value:,.2f}"
            # Inverte separadores: de "1,234.56" para "1.234,56"
            formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
            return f"R$ {formatted}"

        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(query)
                rows = cursor.fetchall()
                if not rows:
                    return "Nenhum resultado retornado."

                # Obter os nomes das colunas
                columns = [desc[0] for desc in cursor.description]
                if "cod_siafi" not in columns:
                    warning = "AVISO: A consulta não retornou a coluna 'cod_siafi'.\n"
                else:
                    warning = ""
                cod_siafi_index = columns.index("cod_siafi") if "cod_siafi" in columns else None

                result_lines = [warning]
                for row in rows:
                    # Se 'cod_siafi' estiver presente, obtém os dados da organização
                    if cod_siafi_index is not None:
                        cod_siafi = row[cod_siafi_index]
                        cursor.execute(
                            "SELECT sigla_om, nome_om, distrito, uf FROM organizacoes_militares WHERE cod_siafi = ?",
                            (cod_siafi,)
                        )
                        org_result = cursor.fetchone()
                        if org_result:
                            sigla_om, nome_om, distrito, uf = org_result
                            org_info = f"Organização Militar: {sigla_om} - {nome_om} ({distrito}, {uf})"
                        else:
                            org_info = "Organização não encontrada"
                        result_lines.append(org_info)

                    # Formata os dados da linha em pares "coluna: valor", ignorando valores None
                    row_items = []
                    for col, val in zip(columns, row):
                        if val is None:
                            continue
                        # Não formata 'cod_siafi' como moeda
                        if isinstance(val, (int, float)) and col != "cod_siafi":
                            val_str = format_currency(val)
                        else:
                            val_str = str(val)
                        row_items.append(f"{col}: {val_str}")
                    if row_items:
                        row_str = ", ".join(row_items)
                        result_lines.append(row_str)
                    result_lines.append("-" * 40)  # linha separadora
                return "\n".join(result_lines)
        except Exception as e:
            logging.error(f"Erro ao executar a consulta SQL: {e}")
            return f"Erro ao executar a consulta SQL: {e}"

        except Exception as e:
            logging.error(f"Erro ao executar a consulta SQL: {e}")
            return f"Erro ao executar a consulta SQL: {e}"

    # Funções wrapper para exportação (sem argumentos no slot)
    def on_export_xlsx(self):
        if self.last_sql_query:
            result = export_xlsx(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação XLSX", result)
        else:
            QMessageBox.warning(self, "Exportação XLSX", "Consulta SQL não disponível.")

    def on_export_docx(self):
        if self.last_sql_query:
            result = export_docx(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação DOCX", result)
        else:
            QMessageBox.warning(self, "Exportação DOCX", "Consulta SQL não disponível.")

    def on_export_pdf(self):
        if self.last_sql_query:
            result = export_pdf(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação PDF", result)
        else:
            QMessageBox.warning(self, "Exportação PDF", "Consulta SQL não disponível.")
            
def extract_order(query):
    """
    Extrai o modo ("single" ou "list") e o número ordinal desejado da query.
    Se a query contiver dígitos, utiliza-os; senão, busca por palavras ordinais.
    """
    num_match = re.search(r'\b(\d+)\b', query)
    if num_match:
        number = int(num_match.group(1))
        if "maiores" in query or "os" in query:
            return ("list", number)
        else:
            return ("single", number)
    ordinals_map = {
        "primeiro": 1, "primeira": 1,
        "segundo": 2, "segunda": 2,
        "terceiro": 3, "terceira": 3,
        "quarto": 4, "quarta": 4,
        "quinto": 5, "quinta": 5,
        "sexto": 6, "sexta": 6,
        "sétimo": 7, "sétima": 7,
        "oitavo": 8, "oitava": 8,
        "nono": 9, "nona": 9,
        "décimo": 10, "décima": 10,
        "décimo primeiro": 11, "décima primeira": 11,
        "décimo segundo": 12, "décima segunda": 12,
        "décimo terceiro": 13, "décima terceira": 13,
        "décimo quarto": 14, "décima quarta": 14
    }
    for word, num in ordinals_map.items():
        if word in query:
            if "maiores" in query or "os" in query:
                return ("list", num)
            else:
                return ("single", num)
    if "valor" in query:
        return ("single", 1)
    return (None, None)

import os
import sqlite3
import logging
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def export_xlsx(db_path, query, output_filename="export.xlsx"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo XLSX.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para XLSX: {e}")
        return f"Erro: {e}"

    wb = Workbook()
    ws = wb.active
    ws.append(columns)
    for row in rows:
        ws.append(row)
    try:
        wb.save(output_filename)
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao salvar ou abrir o arquivo XLSX: {e}")
        return f"Erro: {e}"
    return f"Arquivo XLSX exportado: {output_filename}"

def export_docx(db_path, query, output_filename="export.docx"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo DOCX.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para DOCX: {e}")
        return f"Erro: {e}"

    document = Document()
    document.add_heading("Exportação de Consulta SQL", level=0)
    
    # Cria uma tabela com cabeçalho
    table = document.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = str(col)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell) if cell is not None else ""
    try:
        document.save(output_filename)
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao salvar ou abrir o arquivo DOCX: {e}")
        return f"Erro: {e}"
    return f"Arquivo DOCX exportado: {output_filename}"

def export_pdf(db_path, query, output_filename="export.pdf"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo PDF.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para PDF: {e}")
        return f"Erro: {e}"

    try:
        c = canvas.Canvas(output_filename, pagesize=letter)
        width, height = letter
        textobject = c.beginText(40, height - 40)
        textobject.setFont("Helvetica", 10)
        
        header_line = " | ".join(str(col) for col in columns)
        textobject.textLine(header_line)
        textobject.textLine("-" * 100)
        for row in rows:
            line = " | ".join(str(cell) if cell is not None else "" for cell in row)
            textobject.textLine(line)
            # Se o texto estiver próximo da parte inferior da página, cria nova página.
            if textobject.getY() < 40:
                c.drawText(textobject)
                c.showPage()
                textobject = c.beginText(40, height - 40)
                textobject.setFont("Helvetica", 10)
        c.drawText(textobject)
        c.save()
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao gerar ou abrir o arquivo PDF: {e}")
        return f"Erro: {e}"
    return f"Arquivo PDF exportado: {output_filename}"
