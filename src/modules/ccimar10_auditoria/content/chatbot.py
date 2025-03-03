from PyQt6.QtWidgets import (
    QFrame, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QLineEdit, QTextEdit,
    QMessageBox, QCheckBox, QGroupBox
)
from paths.base_path import API_KEY
import openai
import sqlite3
import logging
from PyQt6.QtCore import Qt, pyqtSignal
import re
from utils.add_button import add_button_func
from .chatbot_utils.flow_layout import FlowLayout

def create_chatbot(title_text, database_model, icons):
    """Creates and returns the chatbot interface."""
    openai_api_key = API_KEY  # Obtém a API_KEY
    content_frame = QFrame()
    layout = QVBoxLayout(content_frame)
    layout.setContentsMargins(10, 10, 10, 10)
    layout.setSpacing(10)
    # Chama o ChatbotWidget com a ordem correta:
    chatbot_widget = ChatbotWidget(title_text, openai_api_key, database_model, icons)
    layout.addWidget(chatbot_widget)
    return content_frame


class ChatbotWidget(QFrame):
    def __init__(self, title_text, openai_api_key, database_model, icons, parent=None):
        super().__init__(parent)
        self.icons = icons
        self.api_key = openai_api_key or API_KEY
        if not self.api_key:
            raise ValueError("OpenAI API Key is missing. Please check your config.json or environment variables.")
        self.client = openai.OpenAI(api_key=self.api_key)
        self.icons = icons
        # Verifica se o objeto possui o atributo "database_manager"
        if hasattr(database_model, "database_manager"):
            self.db_path = database_model.database_manager.db_path
        else:
            self.db_path = str(database_model)
        self.db_metadata = self.get_database_metadata()
        self.last_sql_query = None  # Armazena a última consulta SQL

        
        self.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #FFFFFF;
            }
            QLineEdit, QTextEdit {
                background-color: #FFFFFF;
                border-radius: 5px;
                padding: 5px;
                font-size: 14px;
            }
            QPushButton {
                background-color: #6272A4;
                color: #FFF;
                font-weight: bold;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #50FA7B;
            }
            QGroupBox {
                color: white;
                font-size: 14px;
            }
                           
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)
        title = QLabel(title_text)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        self.input_field = QLineEdit()
        self.input_field.setPlaceholderText("Faça a sua pergunta...")
        layout.addWidget(self.input_field)

        # Área de seleção para as tabelas (pode-se selecionar mais de uma)
        self.table_selection_group = QGroupBox("Selecionar Tabelas para Análise")
        self.table_selection_layout = FlowLayout()  # Já com margens e espaçamento configurados
        self.table_checkboxes = {}
        for table_name in self.db_metadata.keys():
            checkbox = QCheckBox(table_name)
            checkbox.setStyleSheet("color: white; font-size: 14px;")
            checkbox.setChecked(False)
            self.table_checkboxes[table_name] = checkbox
            self.table_selection_layout.addWidget(checkbox)
        self.table_selection_group.setLayout(self.table_selection_layout)
        layout.addWidget(self.table_selection_group)


        # Layout horizontal para os botões
        button_layout = QHBoxLayout()
        self.submit_button = QPushButton("Banco de Dados")
        self.submit_button.clicked.connect(self.generate_response)
        button_layout.addWidget(self.submit_button)
        
        self.api_button = QPushButton("API Direta")
        self.api_button.clicked.connect(self.generate_direct_api_response)
        button_layout.addWidget(self.api_button)
        layout.addLayout(button_layout)
        
        self.response_output = QTextEdit()
        self.response_output.setReadOnly(True)
        layout.addWidget(self.response_output)

        # Área de exportação: 3 botões (inicialmente desativados)
        self.export_buttons_layout = QHBoxLayout()
        # Use os métodos wrapper on_export_xlsx, on_export_docx, on_export_pdf (definidos abaixo)
        self.xlsx_button = add_button_func("Tabela XLSX", "excel", self.on_export_xlsx, self.export_buttons_layout,
                                        icons=self.icons, tooltip="Exportar tabela para XLSX", button_size=(200, 40))
        self.docx_button = add_button_func("Documento DOCX", "word", self.on_export_docx, self.export_buttons_layout,
                                        icons=self.icons, tooltip="Exportar documento para DOCX", button_size=(200, 40))
        self.pdf_button = add_button_func("Relatório PDF", "pdf", self.on_export_pdf, self.export_buttons_layout,
                                        icons=self.icons, tooltip="Exportar relatório para PDF", button_size=(200, 40))


        self.xlsx_button.setEnabled(False)
        self.docx_button.setEnabled(False)
        self.pdf_button.setEnabled(False)
        layout.addLayout(self.export_buttons_layout)

    def get_database_metadata(self):
        """Lê as tabelas e suas colunas no banco de dados."""
        metadata = {}
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                for table_name in tables:
                    table_name = table_name[0]
                    cursor.execute(f"PRAGMA table_info({table_name})")
                    columns = cursor.fetchall()
                    # Armazena os nomes das colunas com seus tipos, ex.: "cod_siafi (INTEGER)"
                    metadata[table_name] = [f"{col[1]} ({col[2]})" for col in columns]
        except sqlite3.Error as e:
            logging.error(f"Database error: {e}")
            return {}
        return metadata

    def get_max_value_from_table(self, column_name="valor_total_execucao_licitacao"):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                query = f"""
                    SELECT cod_siafi, MAX({column_name}) 
                    FROM criterio_execucao_licitacao
                    WHERE {column_name} IS NOT NULL;
                """
                cursor.execute(query)
                result = cursor.fetchone()
                if not result or result[1] is None:
                    return f"Nenhum valor encontrado na coluna '{column_name}'."
                cod_siafi, max_value = result
                cursor.execute("""
                    SELECT sigla_om, nome_om, distrito, uf 
                    FROM organizacoes_militares 
                    WHERE cod_siafi = ?;
                """, (cod_siafi,))
                org_result = cursor.fetchone()
                if org_result:
                    sigla_om, nome_om, distrito, uf = org_result
                    org_info = f"Organização Militar: {sigla_om} - {nome_om} ({distrito}, {uf})"
                else:
                    org_info = "Nenhuma organização militar encontrada para este cod_siafi."
                return f"O maior valor da coluna '{column_name}' é {max_value} (cod_siafi: {cod_siafi}).\n{org_info}"
        except sqlite3.Error as e:
            logging.error(f"Database error: {e}")
            return "Erro ao acessar o banco de dados."

    def get_value_by_rank_from_table(self, column, rank, mode="single"):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                if mode == "single":
                    query = f"""
                        SELECT cod_siafi, {column}
                        FROM criterio_execucao_licitacao
                        WHERE {column} IS NOT NULL
                        ORDER BY {column} DESC
                        LIMIT 1 OFFSET {rank - 1};
                    """
                    cursor.execute(query)
                    result = cursor.fetchone()
                    if not result or result[1] is None:
                        return f"Nenhum valor encontrado na coluna '{column}' para a posição {rank}."
                    cod_siafi, value = result
                    cursor.execute("""
                        SELECT sigla_om, nome_om, distrito, uf 
                        FROM organizacoes_militares 
                        WHERE cod_siafi = ?;
                    """, (cod_siafi,))
                    org_result = cursor.fetchone()
                    if org_result:
                        sigla_om, nome_om, distrito, uf = org_result
                        org_info = f"Organização Militar: {sigla_om} - {nome_om} ({distrito}, {uf})"
                    else:
                        org_info = "Nenhuma organização militar encontrada para este cod_siafi."
                    return f"O {rank}º valor da coluna '{column}' é {value} (cod_siafi: {cod_siafi}).\n{org_info}"
                elif mode == "list":
                    query = f"""
                        SELECT cod_siafi, {column}
                        FROM criterio_execucao_licitacao
                        WHERE {column} IS NOT NULL
                        ORDER BY {column} DESC
                        LIMIT {rank};
                    """
                    cursor.execute(query)
                    results = cursor.fetchall()
                    if not results:
                        return f"Nenhum valor encontrado na coluna '{column}'."
                    response_lines = []
                    for idx, row in enumerate(results, start=1):
                        cod_siafi, value = row
                        cursor.execute("""
                            SELECT sigla_om, nome_om, distrito, uf 
                            FROM organizacoes_militares 
                            WHERE cod_siafi = ?;
                        """, (cod_siafi,))
                        org_result = cursor.fetchone()
                        if org_result:
                            sigla_om, nome_om, distrito, uf = org_result
                            org_info = f"{sigla_om} - {nome_om} ({distrito}, {uf})"
                        else:
                            org_info = "Organização não encontrada"
                        response_lines.append(f"{idx}º: {value} (cod_siafi: {cod_siafi}) - {org_info}")
                    return "\n".join(response_lines)
        except sqlite3.Error as e:
            logging.error(f"Database error: {e}")
            return "Erro ao acessar o banco de dados."

    def generate_response(self):
        """Gera resposta com base na pergunta do usuário, identificando a coluna e o ranking desejado."""
        user_question = self.input_field.text().strip().lower()
        if not user_question:
            QMessageBox.warning(self, "Warning", "Por favor, insira uma pergunta.")
            return

        specific_keywords = {
            "valor_pregao_eletronico": ["pregao eletronico", "pregão eletrônico", "pregao eletRônico"],
            "valor_concorrencia": ["concorrencia", "concorrência", "concorrnca", "concorrrencia"],
            "valor_credenciamento": ["credenciamento", "credenciamnto", "credenciamento"]
        }
        other_mapping = {
            "valor_convite": ["convite"],
            "valor_tomada_preco": ["tomada de preco", "tomada de preço"],
            "valor_dispensa": ["dispensa"],
            "valor_inexigibilidade": ["inexigibilidade"],
            "valor_nao_se_aplica": ["nao se aplica", "não se aplica"],
            "valor_suprimento_fundos": ["suprimento"],
            "valor_regime_diferenciado": ["regime diferenciado"],
            "valor_cons": ["cons"]
        }
        column = None
        for col, variants in specific_keywords.items():
            if any(variant in user_question for variant in variants):
                column = col
                break
        if not column:
            for col, variants in other_mapping.items():
                if any(variant in user_question for variant in variants):
                    column = col
                    break
        if not column and "valor" in user_question:
            column = "valor_total_execucao_licitacao"

        mode, order_num = extract_order(user_question)
        if column and order_num is not None:
            response_text = self.get_value_by_rank_from_table(column, order_num, mode)
            self.response_output.setText(response_text)
            return

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": f"Você é um assistente especializado em análise de dados do banco SQLite. Analise a tabela {self.db_metadata} e responda diretamente o que foi perguntado com os dados obtidos da tabela. Se for necessário informe também o código sql para obter os dados"},
                    {"role": "user", "content": f"Pergunta: {self.input_field.text().strip()}"}
                ]
            )
            ai_response = response.choices[0].message.content
        except Exception as e:
            logging.error(f"Error connecting to OpenAI: {e}")
            ai_response = "Erro ao acessar a API do OpenAI."
        self.response_output.setText(ai_response)

    def generate_direct_api_response(self):
        """
        Envia a pergunta diretamente à API, incluindo os metadados das tabelas selecionadas.
        A resposta da API deverá ser uma consulta SQL que será executada no banco de dados e
        o resultado será exibido logo abaixo da resposta do ChatGPT.
        A consulta SQL DEVE SEMPRE incluir a coluna 'cod_siafi'.
        """
        user_question = self.input_field.text().strip()
        if not user_question:
            QMessageBox.warning(self, "Warning", "Por favor, insira uma pergunta.")
            return
        
        selected_tables = []
        for table, checkbox in self.table_checkboxes.items():
            if checkbox.isChecked():
                columns_info = ", ".join(self.db_metadata.get(table, []))
                selected_tables.append(f"Tabela: {table}\nColunas: {columns_info}")
        if not selected_tables:
            QMessageBox.warning(self, "Warning", "Por favor, selecione ao menos uma tabela para análise.")
            return
        tables_info = "\n\n".join(selected_tables)
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Você é um assistente especializado em análise de dados do banco SQLite. "
                            "Considere os metadados das tabelas abaixo e responda à pergunta retornando SOMENTE uma consulta SQL, "
                            "a qual deverá ser utilizada para consultar o banco de dados e retornar os valores encontrados. "
                            "A consulta SQL DEVE SEMPRE incluir a coluna 'cod_siafi' para facilitar o rastreio.\n\n"
                            f"{tables_info}"
                        )
                    },
                    {"role": "user", "content": f"Pergunta: {user_question}"}
                ]
            )
            ai_response = response.choices[0].message.content
        except Exception as e:
            logging.error(f"Error connecting to OpenAI: {e}")
            self.response_output.setText("Erro ao acessar a API do OpenAI.")
            return

        output_text = f"Consulta SQL gerada:\n{ai_response}\n\n"
        sql_query = ""
        if "```" in ai_response:
            parts = ai_response.split("```")
            if len(parts) >= 2:
                sql_query = parts[1].strip()
                if sql_query.lower().startswith("sql"):
                    sql_query = sql_query[3:].strip()
        else:
            sql_query = ai_response.strip()

        if sql_query:
            self.last_sql_query = sql_query  # Armazena a consulta para exportação
            logging.info(f"Consulta SQL armazenada: {self.last_sql_query}")
            result_text = self.execute_sql_query(sql_query)
            output_text += f"Resultados da Consulta:\n{result_text}"

            self.xlsx_button.setEnabled(True)
            self.docx_button.setEnabled(True)
            self.pdf_button.setEnabled(True)
        else:
            output_text += "Não foi possível extrair a consulta SQL."
            self.xlsx_button.setEnabled(False)
            self.docx_button.setEnabled(False)
            self.pdf_button.setEnabled(False)

        self.response_output.setText(output_text)

    def execute_sql_query(self, query):
        """
        Executa a consulta SQL no banco de dados e retorna os resultados formatados.
        Se a consulta retornar a coluna 'cod_siafi', para cada linha é realizada uma consulta
        na tabela 'organizacoes_militares' para obter os dados da organização correspondente.
        Valores numéricos são convertidos para o formato BRL (R$ 1.234,56) e colunas com valor None são ignoradas.
        """
        def format_currency(value):
            # Formata valores numéricos para o padrão BRL: 1234.56 -> "R$ 1.234,56"
            formatted = f"{value:,.2f}"
            # Inverte separadores: de "1,234.56" para "1.234,56"
            formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
            return f"R$ {formatted}"

        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(query)
                rows = cursor.fetchall()
                if not rows:
                    return "Nenhum resultado retornado."

                # Obter os nomes das colunas
                columns = [desc[0] for desc in cursor.description]
                if "cod_siafi" not in columns:
                    warning = "AVISO: A consulta não retornou a coluna 'cod_siafi'.\n"
                else:
                    warning = ""
                cod_siafi_index = columns.index("cod_siafi") if "cod_siafi" in columns else None

                result_lines = [warning]
                for row in rows:
                    # Se 'cod_siafi' estiver presente, obtém os dados da organização
                    if cod_siafi_index is not None:
                        cod_siafi = row[cod_siafi_index]
                        cursor.execute(
                            "SELECT sigla_om, nome_om, distrito, uf FROM organizacoes_militares WHERE cod_siafi = ?",
                            (cod_siafi,)
                        )
                        org_result = cursor.fetchone()
                        if org_result:
                            sigla_om, nome_om, distrito, uf = org_result
                            org_info = f"Organização Militar: {sigla_om} - {nome_om} ({distrito}, {uf})"
                        else:
                            org_info = "Organização não encontrada"
                        result_lines.append(org_info)

                    # Formata os dados da linha em pares "coluna: valor", ignorando valores None
                    row_items = []
                    for col, val in zip(columns, row):
                        if val is None:
                            continue
                        # Não formata 'cod_siafi' como moeda
                        if isinstance(val, (int, float)) and col != "cod_siafi":
                            val_str = format_currency(val)
                        else:
                            val_str = str(val)
                        row_items.append(f"{col}: {val_str}")
                    if row_items:
                        row_str = ", ".join(row_items)
                        result_lines.append(row_str)
                    result_lines.append("-" * 40)  # linha separadora
                return "\n".join(result_lines)
        except Exception as e:
            logging.error(f"Erro ao executar a consulta SQL: {e}")
            return f"Erro ao executar a consulta SQL: {e}"

        except Exception as e:
            logging.error(f"Erro ao executar a consulta SQL: {e}")
            return f"Erro ao executar a consulta SQL: {e}"

    # Funções wrapper para exportação (sem argumentos no slot)
    def on_export_xlsx(self):
        if self.last_sql_query:
            result = export_xlsx(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação XLSX", result)
        else:
            QMessageBox.warning(self, "Exportação XLSX", "Consulta SQL não disponível.")

    def on_export_docx(self):
        if self.last_sql_query:
            result = export_docx(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação DOCX", result)
        else:
            QMessageBox.warning(self, "Exportação DOCX", "Consulta SQL não disponível.")

    def on_export_pdf(self):
        if self.last_sql_query:
            result = export_pdf(self.db_path, self.last_sql_query)
            QMessageBox.information(self, "Exportação PDF", result)
        else:
            QMessageBox.warning(self, "Exportação PDF", "Consulta SQL não disponível.")
            

def extract_order(query):
    """
    Extrai o modo ("single" ou "list") e o número ordinal desejado da query.
    Se a query contiver dígitos, utiliza-os; senão, busca por palavras ordinais.
    """
    num_match = re.search(r'\b(\d+)\b', query)
    if num_match:
        number = int(num_match.group(1))
        if "maiores" in query or "os" in query:
            return ("list", number)
        else:
            return ("single", number)
    ordinals_map = {
        "primeiro": 1, "primeira": 1,
        "segundo": 2, "segunda": 2,
        "terceiro": 3, "terceira": 3,
        "quarto": 4, "quarta": 4,
        "quinto": 5, "quinta": 5,
        "sexto": 6, "sexta": 6,
        "sétimo": 7, "sétima": 7, "setimo": 7, "setima": 7,
        "oitavo": 8, "oitava": 8,
        "nono": 9, "nona": 9,
        "décimo": 10, "décima": 10, "decimo": 10, "decima": 10,
        "décimo primeiro": 11, "décima primeira": 11, "decimo primeiro": 11, "decima primeira": 11,
        "décimo segundo": 12, "décima segunda": 12, "decimo segundo": 12, "decima segunda": 12,
        "décimo terceiro": 13, "décima terceira": 13, "decimo terceiro": 13, "decima terceira": 13,
        "décimo quarto": 14, "décima quarta": 14, "decimo quarto": 14, "decima quarta": 14,
        "décimo quinto": 15, "décima quinta": 15, "decimo quinto": 15, "decima quinta": 15,
        "décimo sexto": 16, "décima sexta": 16, "decimo sexto": 16, "decima sexta": 16,
        "décimo sétimo": 17, "décima sétima": 17, "decimo sétimo": 17, "decima sétima": 17,
        "décimo oitavo": 18, "décima oitava": 18, "decimo oitavo": 18, "decima oitava": 18,
        "décimo nono": 19, "décima nona": 19, "decimo nono": 19, "decima nona": 19,
        "vigésimo": 20, "vigésima": 20, "vigesimo": 20, "vigesima": 20,
        "vigésimo primeiro": 21, "vigésima primeira": 21, "vigesimo primeiro": 21, "vigesima primeira": 21,
        "vigésimo segundo": 22, "vigésima segunda": 22, "vigesimo segundo": 22, "vigesima segunda": 22,
        "vigésimo terceiro": 23, "vigésima terceira": 23, "vigesimo terceiro": 23, "vigesima terceira": 23,
        "vigésimo quarto": 24, "vigésima quarta": 24, "vigesimo quarto": 24, "vigesima quarta": 24,
        "vigésimo quinto": 25, "vigésima quinta": 25, "vigesimo quinto": 25, "vigesima quinta": 25,
        "vigésimo sexto": 26, "vigésima sexta": 26, "vigesimo sexto": 26, "vigesima sexta": 26,
        "vigésimo sétimo": 27, "vigésima sétima": 27, "vigesimo sétimo": 27, "vigesima sétima": 27,
        "vigésimo oitavo": 28, "vigésima oitava": 28, "vigesimo oitavo": 28, "vigesima oitava": 28,
        "vigésimo nono": 29, "vigésima nona": 29, "vigesimo nono": 29, "vigesima nona": 29,
        "trigésimo": 30, "trigésima": 30, "trigesimo": 30, "trigesima": 30,
        "quadragesimo": 40, "quadragesima": 40,
    }
    for word, num in ordinals_map.items():
        if word in query:
            if "maiores" in query or "os" in query:
                return ("list", num)
            else:
                return ("single", num)
    if "valor" in query:
        return ("single", 1)
    return (None, None)

import os
import sqlite3
import logging
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def export_xlsx(db_path, query, output_filename="export.xlsx"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo XLSX.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para XLSX: {e}")
        return f"Erro: {e}"

    wb = Workbook()
    ws = wb.active
    ws.append(columns)
    for row in rows:
        ws.append(row)
    try:
        wb.save(output_filename)
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao salvar ou abrir o arquivo XLSX: {e}")
        return f"Erro: {e}"
    return f"Arquivo XLSX exportado: {output_filename}"

def export_docx(db_path, query, output_filename="export.docx"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo DOCX.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para DOCX: {e}")
        return f"Erro: {e}"

    document = Document()
    document.add_heading("Exportação de Consulta SQL", level=0)
    
    # Cria uma tabela com cabeçalho
    table = document.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = str(col)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell) if cell is not None else ""
    try:
        document.save(output_filename)
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao salvar ou abrir o arquivo DOCX: {e}")
        return f"Erro: {e}"
    return f"Arquivo DOCX exportado: {output_filename}"

def export_pdf(db_path, query, output_filename="export.pdf"):
    """
    Executa a consulta SQL no banco de dados e exporta os resultados para um arquivo PDF.
    Em seguida, abre o arquivo gerado.
    """
    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
    except Exception as e:
        logging.error(f"Erro ao executar a consulta SQL para PDF: {e}")
        return f"Erro: {e}"

    try:
        c = canvas.Canvas(output_filename, pagesize=letter)
        width, height = letter
        textobject = c.beginText(40, height - 40)
        textobject.setFont("Helvetica", 10)
        
        header_line = " | ".join(str(col) for col in columns)
        textobject.textLine(header_line)
        textobject.textLine("-" * 100)
        for row in rows:
            line = " | ".join(str(cell) if cell is not None else "" for cell in row)
            textobject.textLine(line)
            # Se o texto estiver próximo da parte inferior da página, cria nova página.
            if textobject.getY() < 40:
                c.drawText(textobject)
                c.showPage()
                textobject = c.beginText(40, height - 40)
                textobject.setFont("Helvetica", 10)
        c.drawText(textobject)
        c.save()
        os.startfile(output_filename)
    except Exception as e:
        logging.error(f"Erro ao gerar ou abrir o arquivo PDF: {e}")
        return f"Erro: {e}"
    return f"Arquivo PDF exportado: {output_filename}"
